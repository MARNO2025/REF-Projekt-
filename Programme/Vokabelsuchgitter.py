import random
import string
import os
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

# ----------------------------
# Suchnetz erstellen
# ----------------------------
def create_wordgrid(words_list, translations_list, size=20):
    words_list = [w.replace(" ", "") for w in words_list]
    random.shuffle(words_list)
    grid = [['' for _ in range(size)] for _ in range(size)]
    placed_words = []

    for word, translation in zip(words_list, translations_list):
        word_upper = word.upper()
        placed = False
        attempts = 0

        while not placed and attempts < 200:
            attempts += 1
            row = random.randint(0, size-1)
            col = random.randint(0, size-1)
            direction = random.choice(['H', 'V'])

            if direction == 'H' and col + len(word_upper) <= size:
                if all(grid[row][col+i] in ('', word_upper[i]) for i in range(len(word_upper))):
                    for i in range(len(word_upper)):
                        grid[row][col+i] = word_upper[i]
                    placed = True
            elif direction == 'V' and row + len(word_upper) <= size:
                if all(grid[row+i][col] in ('', word_upper[i]) for i in range(len(word_upper))):
                    for i in range(len(word_upper)):
                        grid[row+i][col] = word_upper[i]
                    placed = True

        if placed:
            placed_words.append((word, translation))

    # Leere Felder zufällig füllen
    for i in range(size):
        for j in range(size):
            if grid[i][j] == '':
                grid[i][j] = random.choice(string.ascii_uppercase)

    return grid, placed_words

# ----------------------------
# Word-Dokument erstellen
# ----------------------------
def create_word_doc(grid, placed_words, template_path):
    doc = Document(template_path)
    doc.add_heading("Cherche le vocabulaire", level=1)

    # 🔲 Suchnetz-Tabelle
    table_grid = doc.add_table(rows=len(grid), cols=len(grid[0]))
    table_grid.style = 'Table Grid'

    for i, row in enumerate(grid):
        for j, letter in enumerate(row):
            cell = table_grid.cell(i, j)
            p = cell.paragraphs[0]

            # run erzeugen, falls keiner existiert
            if not p.runs:
                r = p.add_run(letter)
            else:
                r = p.runs[0]
                r.text = letter

            r.font.name = 'Courier New'
            r.font.size = Pt(12)

    doc.add_paragraph("\n1. Mets les bons mots :\n")

    # 📝 Tabelle für Schülerantworten
    table_answer = doc.add_table(rows=1, cols=2)
    table_answer.style = 'Table Grid'

    # Zelle 0,0 fett
    cell = table_answer.cell(0, 0)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run("mot")
    run.bold = True

    # Zelle 0,1 fett
    cell = table_answer.cell(0, 1)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run("traduction")
    run.bold = True


    # 🔧 Zeilenhöhe setzen
    def set_row_height(row, height_cm):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_cm * 567)))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

# Kopfzeile
    set_row_height(table_answer.rows[0], 1)

# Schülerantwort-Zeilen
    for word, translation in placed_words:
        row = table_answer.add_row()
        cells = row.cells
        cells[0].text = ""           # Schülerantwort
        cells[1].text = translation  # Übersetzung
        set_row_height(row, 1)       # 👉 jetzt jede Zeile gleich hoch machen

    # 📄 Word zurückgeben
    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)
    return word_stream

# ----------------------------
# Hauptfunktion für app.py
# ----------------------------
def run_Vokabelsuchgitter(words_with_translations, template_path):
    """
    Nimmt eine Liste von Tupeln (word, translation) und erstellt ein Word-Dokument als BytesIO.
    """
    words = [w for w, _ in words_with_translations]
    translations = [t for _, t in words_with_translations]

    grid, placed_words = create_wordgrid(words, translations, size=20)
    word_file = create_word_doc(grid, placed_words, template_path=template_path)
    return word_file
