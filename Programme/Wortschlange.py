import random
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from io import BytesIO

def run_Wortschlange(words_with_translations, template_path=None, font_size=14):
    doc = Document(template_path)

    def set_row_height(row, height_cm):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_cm * 567)))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    def add_wortschlange_table(text):
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(font_size)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        doc.add_paragraph("\n")
        return table

    def add_schueler_tabelle():
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["mot", "traduction"]
        for col_idx, text in enumerate(headers):
            cell = table.cell(0, col_idx)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.add_run(text)
            run.bold = True
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            run.font.size = Pt(font_size)
        set_row_height(table.rows[0], 1.2)

        # Zeilen für jede Aufgabe
        for word, translation in words_with_translations:
            row = table.add_row()
            set_row_height(row, 1.2)
            row_cells = row.cells
            # Wort
            p1 = row_cells[0].paragraphs[0]
            run1 = p1.add_run('')
            run1.font.name = 'Arial'
            run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            run1.font.size = Pt(font_size)
            # Übersetzung
            p2 = row_cells[1].paragraphs[0]
            run2 = p2.add_run('')
            run2.font.name = 'Arial'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            run2.font.size = Pt(font_size)
        doc.add_paragraph("\n")
        return table

    # Überschrift für erstes AB
    doc.add_heading("Cherche le mot", level=1)

    # Wortschlange mit Artikeln
    wörter = [word for word, _ in words_with_translations]
    random.shuffle(wörter)
    wortschlange_mit = ''.join([word.lower().replace(" ", "") for word in wörter])
    add_wortschlange_table(wortschlange_mit)
    add_schueler_tabelle()

    # Überschrift für zweites AB
    doc.add_heading("Cherche le mot", level=1)

    # Wortschlange ohne Artikel
    artikel = ["le ", "la ", "l'", "les ", "un ", "une ", "des "]
    wörter_ohne = []
    for word in wörter:
        for art in artikel:
            if word.lower().startswith(art):
                word = word[len(art):]
        wörter_ohne.append(word)
    wortschlange_ohne = ''.join([word.lower().replace(" ", "") for word in wörter_ohne])
    add_wortschlange_table(wortschlange_ohne)
    add_schueler_tabelle()

    # Dokument speichern
    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)
    return word_stream
