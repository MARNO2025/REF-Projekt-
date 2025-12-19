from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import random

def run_konjugationstabelle(words, num_rows, template_path):
    """
    Erstellt eine Konjugationstabelle mit zufälligen Personalpronomen und gibt
    das Word-Dokument als BytesIO zurück.
    
    :param words: Liste von Verben
    :param num_rows: Anzahl der Zeilen
    :param template_path: Pfad zur Word-Vorlage
    :return: BytesIO Objekt des Word-Dokuments
    """
    
    dokument = Document(template_path)

    # Überschrift
    ueberschrift = dokument.add_heading('Test de conjugaison', level=0)
    dokument.add_paragraph('')
    run = ueberschrift.runs[0]
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'
    ueberschrift.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Funktion für Zeilenhöhe
    def set_zeilenhoehe(row, height_cm):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_cm * 567)))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    # Tabelle mit 4 Spalten
    table = dokument.add_table(rows=1, cols=4)
    table.style = 'Tabellenraster'

    # Kopfzeile
    headers = ["Verbe", "Présent", "Passé Composé", "En allemand"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(14)

    personalpronomen = ["je", "tu", "il", "elle", "nous", "vous", "ils", "elles"]

    # Zeilen zufällig füllen
    for _ in range(num_rows):
        verb = random.choice(words)
        pronomen = random.choice(personalpronomen)
        row = table.add_row().cells
        row[0].text = verb
        row[1].text = pronomen
        row[2].text = pronomen 
        row[3].text = "" 

    # Zeilenhöhe setzen
    for row in table.rows:
        set_zeilenhoehe(row, 1)

    # Word-Dokument als BytesIO zurückgeben
    word_stream = BytesIO()
    dokument.save(word_stream)
    word_stream.seek(0)
    return word_stream
