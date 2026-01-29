import random
import os
from docx import Document
from docx.shared import Pt
from io import BytesIO
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def run_Rätsel(words_with_translations, template_path=None, font_size=12,
               heading="mystère", instructions="1. Cherche les mots", max_words=None):

    def buchstaben_mischen(wort):
        if len(wort) <= 1:
            return wort
        buchstaben = list(wort.replace(" ", ""))
        shuffled = wort
        while shuffled == wort:
            random.shuffle(buchstaben)
            shuffled = ''.join(buchstaben)
        return shuffled

    def buchstaben_mit_ersten(wort):
        wort_clean = wort.replace(" ", "")
        if len(wort_clean) <= 1:
            return wort_clean
        erster = wort_clean[0]
        rest = list(wort_clean[1:])
        random.shuffle(rest)
        return erster + ''.join(rest)

    # Dokument laden oder neu erstellen
    doc = Document(template_path)

    # -------------------------
    # SEITE 1 – normales Rätsel
    # -------------------------
    doc.add_heading(heading, level=1)
    doc.add_paragraph(instructions)

    random.shuffle(words_with_translations)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    def set_row_height(row, height_cm):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_cm * 567)))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    headers = ["mystère", "mot", "traduction"]
    for col_idx, text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    set_row_height(table.rows[0], 1.2)

    for word, translation in words_with_translations:
        row = table.add_row()
        set_row_height(row, 1.2)
        row_cells = row.cells

        run0 = row_cells[0].paragraphs[0].add_run(buchstaben_mischen(word))
        run0.font.size = Pt(font_size)

        row_cells[1].paragraphs[0].add_run("").font.size = Pt(font_size)
        row_cells[2].paragraphs[0].add_run("").font.size = Pt(font_size)

    # -------------------------
    # SEITE 2 – erster Buchstabe + gemischter Rest
    # -------------------------
    doc.add_page_break()
    doc.add_heading(heading + " – aide", level=1)
    doc.add_paragraph("1. Trouve les mots (première lettre donnée)")

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'

    for col_idx, text in enumerate(headers):
        cell = table2.cell(0, col_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
    set_row_height(table2.rows[0], 1.2)

    for word, translation in words_with_translations:
        row = table2.add_row()
        set_row_height(row, 1.2)
        row_cells = row.cells

        run0 = row_cells[0].paragraphs[0].add_run(buchstaben_mit_ersten(word))
        run0.font.size = Pt(font_size)

        row_cells[1].paragraphs[0].add_run("").font.size = Pt(font_size)
        row_cells[2].paragraphs[0].add_run("").font.size = Pt(font_size)

    # Dokument zurückgeben
    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)
    return word_stream
