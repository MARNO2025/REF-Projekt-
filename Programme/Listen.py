from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def Vokabellisten(words_with_translation, template_path=None, font_size=14):
    doc = Document(template_path) if template_path else Document()

    # Überschrift
    ueberschrift = doc.add_heading('vocabulaire', level=0)
    ueberschrift.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = ueberschrift.runs[0]
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'

    doc.add_paragraph('')

    def set_zeilenhoehe(row, height_cm):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_cm * 567)))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    # Tabelle
    tabelle = doc.add_table(rows=1, cols=2)
    tabelle.style = 'Table Grid'

    kopfzeile = tabelle.rows[0].cells
    kopfzeile[0].text = "Wort"
    kopfzeile[1].text = "Übersetzung"

    for zelle in kopfzeile:
        p = zelle.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.runs[0]
        r.bold = True
        r.font.size = Pt(font_size)

    for eintrag in words_with_translation:
        zeile = tabelle.add_row().cells
        zeile[0].text = eintrag[0]  # Wort
        zeile[1].text = eintrag[1]  # Zusatzfeld leer
    
    for row in tabelle.rows:
        set_zeilenhoehe(row, 1)

    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)
    return word_stream
