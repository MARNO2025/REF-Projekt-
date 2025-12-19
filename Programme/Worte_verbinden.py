from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import random

def Worte_zuordnen(words_with_translation, template_path=None, font_size=14):
    # Dokument laden oder neu erstellen
    doc = Document(template_path) 

    # Überschrift
    ueberschrift = doc.add_heading('vocabulaire', level=0)
    ueberschrift.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = ueberschrift.runs[0]
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'

    doc.add_paragraph('')

    # Funktion für Zeilenhöhe
    def set_zeilenhoehe(row, height_cm):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_cm * 567)))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    # Tabelle erstellen
    tabelle = doc.add_table(rows=1, cols=2)
    tabelle.style = 'Table Grid'

    # Kopfzeile
    kopfzeile = tabelle.rows[0].cells
    kopfzeile[0].text = "Wort"
    kopfzeile[1].text = "Übersetzung"

    for zelle in kopfzeile:
        p = zelle.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.runs[0]
        r.bold = True
        r.font.size = Pt(font_size)

    # Übersetzungen mischen
    translations = [w[1] for w in words_with_translation]
    random.shuffle(translations)

    # Tabelleninhalt
    for (wort, _), uebersetzung in zip(words_with_translation, translations):
        zeile = tabelle.add_row().cells

        zeile[0].text = "o " + wort
        zeile[1].text = "o " + uebersetzung

        # Schriftgröße setzen
        for zelle in zeile:
            for run in zelle.paragraphs[0].runs:
                run.font.size = Pt(font_size)

    # Zeilenhöhe setzen
    for row in tabelle.rows:
        set_zeilenhoehe(row, 1)

    # Dokument speichern
    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)
    return word_stream
