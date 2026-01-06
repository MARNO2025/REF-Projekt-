from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import random


def run_Unterstriche_Konjugationen(
    words_data,
    num_rows,
    selected_time_1,
    selected_time_2,
    template_path
):
    dokument = Document(template_path)

    # =========================
    # Überschrift Seite 1
    # =========================
    ueberschrift = dokument.add_heading('Test de conjugaison', level=0)
    run = ueberschrift.runs[0]
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial'
    ueberschrift.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dokument.add_paragraph('')

    # =========================
    # Hilfsfunktionen
    # =========================
    def set_zeilenhoehe(row, height_cm):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_cm * 567)))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    def set_spaltenbreiten(table, widths):
        table.autofit = False
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

    # ---------- nur Unterstriche ----------
    def underline_form(form):
        parts = form.split()
        return "   ".join("_ " * len(part) for part in parts).strip()

    # ---------- erster Buchstabe + Unterstriche ----------
    def underline_first_buchstabe(form):
        parts = form.split()
        result = []
        for part in parts:
            if part:
                result.append(part[0] + " " + "_ " * (len(part) - 1))
        return "   ".join(p.strip() for p in result)

    # ---------- französische Elision (immer nach Originalform!) ----------
    def apply_french_elision(pronoun, display_text, original_form):
        vowels = (
            "a", "e", "i", "o", "u", "h",
            "â", "ê", "î", "ô", "û",
            "é", "è", "ë", "ï"
        )

        first_real_char = original_form.strip().lower()[:1]

        if pronoun == "je" and first_real_char in vowels:
            return "j’" + display_text

        return f"{pronoun} {display_text}"

    # =========================
    # Tabelle erstellen
    # =========================
    def create_table(mode="underline"):
        """
        mode:
        - underline     → Unterstriche
        - pronoun       → nur Pronomen
        - first_letter  → erster Buchstabe + Unterstriche
        """
        table = dokument.add_table(rows=1, cols=4)
        table.style = 'Tabellenraster'
        table.autofit = False

        headers = ["Verbe", selected_time_1, selected_time_2, "En allemand"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.bold = True
            run.font.size = Pt(12)

        for entry in exercises:
            row = table.add_row().cells
            row[0].text = entry["verb"]

            if mode == "underline":
                row[1].text = apply_french_elision(
                    entry["p1"],
                    underline_form(entry["f1"]),
                    entry["f1"]
                )
                row[2].text = apply_french_elision(
                    entry["p2"],
                    underline_form(entry["f2"]),
                    entry["f2"]
                )

            elif mode == "pronoun":
                # Elision auch hier korrekt anzeigen
                row[1].text = "j’" if entry["p1"] == "je" and entry["f1"][0].lower() in "aeiouh" else entry["p1"]
                row[2].text = "j’" if entry["p2"] == "je" and entry["f2"][0].lower() in "aeiouh" else entry["p2"]

            elif mode == "first_letter":
                row[1].text = apply_french_elision(
                    entry["p1"],
                    underline_first_buchstabe(entry["f1"]),
                    entry["f1"]
                )
                row[2].text = apply_french_elision(
                    entry["p2"],
                    underline_first_buchstabe(entry["f2"]),
                    entry["f2"]
                )

            row[3].text = ""

        widths = [Cm(3), Cm(5.522), Cm(5.522), Cm(4)]
        set_spaltenbreiten(table, widths)

        for r in table.rows:
            set_zeilenhoehe(r, 1)

    # =========================
    # Aufgaben EINMAL erzeugen
    # =========================
    personalpronomen = ["je", "tu", "il", "elle", "on", "nous", "vous", "ils", "elles"]
    infinitives = list(words_data.keys())

    exercises = []

    for _ in range(num_rows):
        verb = random.choice(infinitives)

        p1 = random.choice(personalpronomen)
        f1 = words_data[verb].get(selected_time_1, {}).get(p1, verb)

        p2 = random.choice(personalpronomen)
        f2 = words_data[verb].get(selected_time_2, {}).get(p2, verb)

        exercises.append({
            "verb": verb,
            "p1": p1,
            "f1": f1,
            "p2": p2,
            "f2": f2
        })

    # =========================
    # Seite 1: Unterstriche
    # =========================
    create_table(mode="underline")

    # =========================
    # Seite 2: nur Pronomen
    # =========================
    dokument.add_page_break()
    u3 = dokument.add_heading('Test de conjugaison – première lettre', level=0)
    u3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    create_table(mode="first_letter")

    # =========================
    # Seite 3: erster Buchstabe
    # =========================
    dokument.add_page_break()
    u2 = dokument.add_heading('Test de conjugaison', level=0)
    u2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    create_table(mode="pronoun")

    # =========================
    # Dokument zurückgeben
    # =========================
    word_stream = BytesIO()
    dokument.save(word_stream)
    word_stream.seek(0)
    return word_stream

