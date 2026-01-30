def generate_worksheets_streamlit(
    text: str,
    vocab_json: list,
    output_prefix: str = "fiche de travail",
    selected_modules: list = [1, 2, 3],
    template_path: str = None
):
    import re
    import random
    from io import BytesIO
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Vorlage laden oder neues Dokument
    if template_path:
        doc = Document(template_path)
    else:
        doc = Document()

    # Vokabeln für Übersetzung
    FR_TO_DE = {v["word"].lower(): v["translation"] for v in vocab_json}

    MODULES = {
        1: "Nur_Unterstriche",
        2: "Erster_Buchstabe",
        3: "Deutsch_in_Box"
    }

    # --- 1. Box einmal erstellen ---
    box = []

    def collect_words(match):
        entry = match.group(1)
        words = entry.split()
        box.append(entry)  # Ursprünglicher Eintrag, für Box
        return match.group(0)  # Text vorerst unverändert

    re.sub(r"\[([^\]]+)\]", collect_words, text)

    # Box einmal mischen
    random.shuffle(box)

    # --- 2. Schleife über Module ---
    for mode in selected_modules:
        mode_name = MODULES[mode]

        def replace(match):
            entry = match.group(1)
            words = entry.split()
            # Text ersetzen je nach Modul
            if mode == 1 or mode == 3:
                return "   ".join(["_ " * len(w) for w in words]).strip()
            elif mode == 2:
                return "   ".join([w[0] + " " + "_ " * (len(w)-1) for w in words]).strip()

        # Lückentext erzeugen
        task_text = re.sub(r"\[([^\]]+)\]", replace, text)

        # --- Überschrift ---
        title = doc.add_heading(f"{output_prefix} – {mode_name.replace('_', ' ')}", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in title.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.font.bold = True

        # --- Wortbox ---
        doc.add_heading("boîte de mots", level=2)
        heading_box = doc.paragraphs[-1]
        heading_box.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in heading_box.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.font.bold = True

        if box:
            # Bei Modul 3: übersetzen
            if mode == 3:
                display_box = [FR_TO_DE.get(w.lower(), w) for w in box]
            else:
                display_box = box

            words_line = " / ".join(display_box)
            p = doc.add_paragraph(words_line)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)

            # Rahmen um die Box
            def set_border(paragraph):
                p_xml = paragraph._p
                pPr = p_xml.get_or_add_pPr()
                pbdr = OxmlElement('w:pBdr')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '6')
                    border.set(qn('w:space'), '1')
                    border.set(qn('w:color'), '000000')
                    pbdr.append(border)
                pPr.append(pbdr)

            set_border(p)

        # --- Aufgabenstellung ---
        task_paragraph = doc.add_paragraph("Ergänze die richtigen Wörter:")
        for run in task_paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(14)

        task_text_paragraph = doc.add_paragraph(task_text)
        for run in task_text_paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)

        # Seitenumbruch zwischen Modulen
        if mode != selected_modules[-1]:
            doc.add_page_break()

    # --- Dokument als BytesIO zurückgeben ---
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

